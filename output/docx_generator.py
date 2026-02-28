"""
DOCX Resume Generator
Generates clean, ATS-friendly Word documents from TailoredResume data.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from engine.resume_tailor import TailoredResume


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        from lxml import etree
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            from lxml import etree
            element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), val.get('space', '0'))


def generate_docx(tailored: TailoredResume, output_path: str):
    """Generate a clean ATS-friendly DOCX resume."""

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    dark = RGBColor(0x1a, 0x1a, 0x1a)
    accent = RGBColor(0x2c, 0x3e, 0x50)
    gray = RGBColor(0x55, 0x55, 0x55)

    personal = tailored.personal

    # === NAME ===
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(personal.get("name", ""))
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.color.rgb = dark
    name_para.paragraph_format.space_after = Pt(2)

    # === CONTACT ===
    contact_parts = []
    if personal.get("email"):
        contact_parts.append(personal["email"])
    for link_name in ["linkedin", "github", "website"]:
        link = personal.get("links", {}).get(link_name, "")
        if link:
            contact_parts.append(link_name.capitalize())
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(" | ".join(contact_parts))
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = gray
    contact_para.paragraph_format.space_after = Pt(4)

    def add_section_header(text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = accent
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(3)
        # Add bottom border
        pPr = para._p.get_or_add_pPr()
        from lxml import etree
        pBdr = etree.SubElement(pPr, qn('w:pBdr'))
        bottom = etree.SubElement(pBdr, qn('w:bottom'))
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'CCCCCC')
        bottom.set(qn('w:space'), '1')
        return para

    def add_exp_header(title_text, date_text):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.columns[0].width = Inches(5.2)
        table.columns[1].width = Inches(1.6)

        left_cell = table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_run = left_para.add_run(title_text)
        left_run.font.size = Pt(10)
        left_run.font.bold = True
        left_run.font.color.rgb = dark
        left_para.paragraph_format.space_after = Pt(0)

        right_cell = table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_run = right_para.add_run(date_text)
        right_run.font.size = Pt(10)
        right_run.font.bold = True
        right_run.font.color.rgb = dark
        right_para.paragraph_format.space_after = Pt(0)

        # Remove table borders
        for cell in [left_cell, right_cell]:
            set_cell_border(cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                start={"val": "none", "sz": "0", "color": "FFFFFF"},
                end={"val": "none", "sz": "0", "color": "FFFFFF"},
            )

    def add_bullet(text):
        para = doc.add_paragraph()
        run = para.add_run(f"● {text}")
        run.font.size = Pt(9.5)
        run.font.color.rgb = dark
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.space_after = Pt(1.5)
        para.paragraph_format.space_before = Pt(0)

    # === EDUCATION ===
    add_section_header("EDUCATION")
    for edu in tailored.education:
        add_exp_header(
            f"{edu.get('institution', '')} | {edu.get('location', '')}",
            f"{edu.get('start_date', '')} – {edu.get('end_date', '')}"
        )
        degree = edu.get("degree", "")
        if edu.get("gpa"):
            degree += f" (GPA: {edu['gpa']})"
        detail = doc.add_paragraph()
        detail_run = detail.add_run(degree)
        detail_run.font.size = Pt(9)
        detail_run.font.italic = True
        detail_run.font.color.rgb = gray
        detail.paragraph_format.space_after = Pt(2)

    # === TECHNICAL SKILLS ===
    add_section_header("TECHNICAL KNOWLEDGE")
    skill_labels = {
        "languages": "Languages",
        "data_engineering": "Data Engineering",
        "cloud_and_tools": "Cloud & Tools",
        "ai_ml": "AI/ML"
    }
    for key, label in skill_labels.items():
        skills = tailored.technical_skills.get(key, [])
        if skills:
            para = doc.add_paragraph()
            label_run = para.add_run(f"{label}: ")
            label_run.font.bold = True
            label_run.font.size = Pt(9.5)
            skills_run = para.add_run(" | ".join(skills))
            skills_run.font.size = Pt(9.5)
            para.paragraph_format.space_after = Pt(2)

    # === WORK EXPERIENCE ===
    if tailored.work_experience:
        add_section_header("WORK EXPERIENCE")
        for exp in tailored.work_experience:
            add_exp_header(
                f"{exp.get('title', '')} | {exp.get('company', '')} | {exp.get('location', '')}",
                f"{exp.get('start_date', '')} – {exp.get('end_date', '')}"
            )
            for bullet in exp.get("bullets", []):
                add_bullet(bullet)

    # === RESEARCH EXPERIENCE ===
    if tailored.research_experience:
        add_section_header("GRADUATE RESEARCH EXPERIENCE")
        for exp in tailored.research_experience:
            add_exp_header(
                f"{exp.get('title', '')} | {exp.get('company', '')} | {exp.get('location', '')}",
                f"{exp.get('start_date', '')} – {exp.get('end_date', '')}"
            )
            for bullet in exp.get("bullets", []):
                add_bullet(bullet)

    # === PROJECTS ===
    if tailored.projects:
        add_section_header("PROJECTS")
        for proj in tailored.projects:
            add_exp_header(
                f"{proj.get('title', '')} | {proj.get('institution', '')}",
                proj.get("date", "")
            )
            for bullet in proj.get("bullets", []):
                add_bullet(bullet)

    # === CERTIFICATIONS ===
    if tailored.certifications:
        add_section_header("CERTIFICATIONS")
        para = doc.add_paragraph()
        run = para.add_run(", ".join(tailored.certifications))
        run.font.size = Pt(9.5)

    doc.save(output_path)
    return output_path
